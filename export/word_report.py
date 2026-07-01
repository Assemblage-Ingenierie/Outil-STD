"""Génération du rapport Word selon la charte Assemblage ingénierie."""
from __future__ import annotations
from pathlib import Path
from io import BytesIO
import tempfile
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import plotly.graph_objects as go


# Chemins assets
ASSETS_DIR = Path(__file__).parent.parent / 'assets'
LOGO_PATH = ASSETS_DIR / 'logos' / 'logo_Ai_rouge_HD.png'
SIGLE_PATH = ASSETS_DIR / 'sigles' / 'sigle_A_rouge_HD.png'

# Couleurs charte
C_ROUGE = RGBColor(0xE3, 0x05, 0x13)
C_VIOLET = RGBColor(0x30, 0x32, 0x3E)
C_GRIS_CLAIR = RGBColor(0xDF, 0xE4, 0xE8)
C_ROUGE_CLAIR = RGBColor(0xF9, 0xE1, 0xE3)
C_GRIS_TC = RGBColor(0xF2, 0xF2, 0xF2)
C_BLANC = RGBColor(0xFF, 0xFF, 0xFF)
C_NOIR70 = RGBColor(0x4D, 0x4D, 0x4D)


def _set_cell_bg(cell, rgb: tuple[int, int, int]):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*rgb))
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _para_rouge(doc: Document, text: str, level: int = 1) -> None:
    sizes = {1: 28, 2: 18, 3: 14}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24 if level == 1 else 16 if level == 2 else 10)
    p.paragraph_format.space_after = Pt(12 if level == 1 else 8 if level == 2 else 6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 14))
    run.font.color.rgb = C_ROUGE if level == 1 else C_VIOLET
    run.font.name = 'Open Sans'


def _table_df(doc: Document, df, index_label: str = "") -> None:
    """Rend un DataFrame en tableau Word stylé charte (en-tête rouge, lignes zébrées).
    L'index est rendu en 1re colonne (libellé `index_label`)."""
    df = df.reset_index()
    if index_label and len(df.columns):
        df = df.rename(columns={df.columns[0]: index_label})
    cols = [str(c) for c in df.columns]
    table = doc.add_table(rows=1 + len(df), cols=len(cols))
    table.style = 'Table Grid'
    for j, col_name in enumerate(cols):
        cell = table.rows[0].cells[j]
        cell.text = col_name
        _set_cell_bg(cell, (0xE3, 0x05, 0x13))
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = C_BLANC
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Open Sans'
    for i, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            if isinstance(val, float) and val != val:      # NaN
                txt = 'NA'
            else:
                txt = str(val)
            cell.text = txt
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(8)
            run.font.name = 'Open Sans'
            if i % 2 == 1:
                _set_cell_bg(cell, (0xF2, 0xF2, 0xF2))
    doc.add_paragraph()


def _fig_to_image(fig: go.Figure) -> BytesIO:
    """Exporte une figure Plotly en PNG dans un buffer."""
    buf = BytesIO()
    fig.write_image(buf, format='png', width=1400, height=700, scale=2)
    buf.seek(0)
    return buf


def generer_rapport(
    variantes: list,
    config: dict,
    seuil_t1: float,
    seuil_t2: float,
    zones_focus: list[str] | None = None,
    zones_comparaison: list[str] | None = None,
    nom_projet: str = "Projet STD",
    methode: str = "givoni",
    df_recap=None,
    df_detail=None,
    seuil_t0: float = 18.0,
    jour_debut: float = 7.0,
    jour_fin: float = 22.0,
    periodes_on: bool = False,
    periodes: list[dict] | None = None,
) -> BytesIO:
    """
    Génère le rapport Word complet selon la trame fixe:
      1. Page de couverture
      2. Synthèse générale (tableau par variante + inconfort jour/nuit)
      3. Focus zones sélectionnées (Givoni + graphiques + humidité)
      4. Comparaison de zones

    Le rapport est toujours rendu en mode clair (légendes foncées lisibles),
    indépendamment du mode sombre éventuellement actif à l'écran.
    """
    from charts.temperature import (
        graphique_heures_depassement,
        graphique_temp_min_moy_max,
        graphique_apports_solaires,
        graphique_apports_par_zone_mensuel,
        graphique_temp_horaire,
        graphique_text_vs_text_op,
        heatmap_temp_jour_heure,
    )
    from charts.givoni import creer_givoni
    from charts.humidite import graphique_hr_horaire, heatmap_hr_jour_heure
    import config.charte as _charte

    # Forcer le mode clair pour l'impression (légendes/annotations foncées,
    # fond blanc) — restauré avant de rendre la main.
    _prev_dark = _charte.is_dark()
    _charte.set_dark_mode(False)

    doc = Document()

    # -- Marges A4 --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(1.5)

    # -- Couverture --
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(1.57))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    titre_p = doc.add_paragraph()
    titre_p.paragraph_format.space_before = Pt(48)
    run_titre = titre_p.add_run(config.get('rapport', {}).get('titre', 'Étude Thermique Dynamique'))
    run_titre.font.size = Pt(24)
    run_titre.font.name = 'Open Sans'
    run_titre.font.color.rgb = C_VIOLET

    sous_titre = doc.add_paragraph()
    run_st = sous_titre.add_run(nom_projet or "")   # add_run crée toujours le run, même si le nom est vide
    run_st.font.size = Pt(16)
    run_st.font.color.rgb = C_ROUGE

    doc.add_page_break()

    # -- 0. Variantes étudiées (récapitulatif + descriptif des améliorations) --
    if (df_recap is not None and not df_recap.empty) or \
       (df_detail is not None and not df_detail.empty):
        _para_rouge(doc, "Variantes étudiées", level=1)

        if df_recap is not None and not df_recap.empty:
            _para_rouge(doc, "Améliorations par variante", level=2)
            # index = améliorations, colonnes = variantes ; cases cochées -> ✕
            rec = df_recap.reset_index()
            cols = list(rec.columns)
            t = doc.add_table(rows=1 + len(rec), cols=len(cols))
            t.style = 'Table Grid'
            for j, c in enumerate(cols):
                cell = t.rows[0].cells[j]
                cell.text = str(c)
                _set_cell_bg(cell, (0xE3, 0x05, 0x13))
                r = cell.paragraphs[0].runs[0]
                r.font.color.rgb = C_BLANC; r.font.bold = True
                r.font.size = Pt(8); r.font.name = 'Open Sans'
            for i, row in rec.iterrows():
                for j, (cname, val) in enumerate(row.items()):
                    cell = t.rows[i + 1].cells[j]
                    if j == 0:
                        txt = str(val)
                    else:
                        txt = "✕" if bool(val) else ""
                    cell.text = txt
                    rr = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(txt)
                    rr.font.size = Pt(8); rr.font.name = 'Open Sans'
                    if j > 0:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if i % 2 == 1:
                        _set_cell_bg(cell, (0xF2, 0xF2, 0xF2))
            doc.add_paragraph()

        if df_detail is not None and not df_detail.empty:
            _para_rouge(doc, "Descriptif des améliorations", level=2)
            dcols = list(df_detail.columns)
            t2 = doc.add_table(rows=1 + len(df_detail), cols=len(dcols))
            t2.style = 'Table Grid'
            for j, c in enumerate(dcols):
                cell = t2.rows[0].cells[j]
                cell.text = str(c)
                _set_cell_bg(cell, (0xE3, 0x05, 0x13))
                r = cell.paragraphs[0].runs[0]
                r.font.color.rgb = C_BLANC; r.font.bold = True
                r.font.size = Pt(8); r.font.name = 'Open Sans'
            for i, row in df_detail.reset_index(drop=True).iterrows():
                for j, val in enumerate(row):
                    cell = t2.rows[i + 1].cells[j]
                    cell.text = str(val) if val is not None else ""
                    rr = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
                    rr.font.size = Pt(8); rr.font.name = 'Open Sans'
                    if i % 2 == 1:
                        _set_cell_bg(cell, (0xF2, 0xF2, 0xF2))
            doc.add_paragraph()

        doc.add_page_break()

    # -- 1. Synthèse générale (1 ligne par variante, niveau bâtiment) --
    _para_rouge(doc, "1. Synthèse générale", level=1)

    import pandas as _pd
    rows = []
    for var in variantes:
        ind = var.indicateurs_batiment(config, methode, seuil_t0=seuil_t0,
                                       seuil_t1=seuil_t1, seuil_t2=seuil_t2)
        rows.append({'Variante': var.nom, **ind})
    df_syn = _pd.DataFrame(rows)

    if df_syn.empty:
        doc.add_paragraph("Aucune donnée disponible.")
    else:
        cols = list(df_syn.columns)
        table = doc.add_table(rows=1 + len(df_syn), cols=len(cols))
        table.style = 'Table Grid'
        for j, col_name in enumerate(cols):
            cell = table.rows[0].cells[j]
            cell.text = col_name
            _set_cell_bg(cell, (0xE3, 0x05, 0x13))
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = C_BLANC
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Open Sans'
        for i, row in df_syn.iterrows():
            for j, (col_name, val) in enumerate(row.items()):
                cell = table.rows[i + 1].cells[j]
                if isinstance(val, float) and val != val:          # NaN
                    txt = 'NA'
                elif col_name.startswith('% hors') and isinstance(val, (int, float)):
                    txt = f"{val:.1f} %"
                else:
                    txt = str(val)
                cell.text = txt
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(txt)
                run.font.size = Pt(8)
                run.font.name = 'Open Sans'
                if i % 2 == 1:
                    _set_cell_bg(cell, (0xF2, 0xF2, 0xF2))
        doc.add_paragraph()

    # -- 1bis. Inconfort par plage horaire (jour / nuit) --
    _para_rouge(doc, "Inconfort par plage horaire (jour / nuit)", level=2)
    rows_dn = []
    for var in variantes:
        dn = var.inconfort_plages_horaires(config, methode, jour_debut, jour_fin)
        rows_dn.append({'Variante': var.nom,
                        **{k: (f"{v:.1f} %" if v == v else 'NA') for k, v in dn.items()}})
    if rows_dn:
        _table_df(doc, _pd.DataFrame(rows_dn).set_index('Variante'), index_label="Variante")
    p_dn = doc.add_paragraph(
        f"Part d'heures d'occupation hors confort, jour [{jour_debut:.0f} h, "
        f"{jour_fin:.0f} h[ / nuit. Agrégat bâtiment pondéré par la surface utile.")
    p_dn.runs[0].font.size = Pt(8)
    p_dn.runs[0].font.color.rgb = C_NOIR70

    # -- 2. Focus zones --
    if zones_focus and variantes:
        doc.add_page_break()
        _para_rouge(doc, "2. Focus zones", level=1)
        # Note chauffe-sous-consigne sur le Givoni : seulement si une variante est chauffée
        note_chauffe = any(v.a_chauffage() for v in variantes)

        for zone in zones_focus:
            _para_rouge(doc, zone, level=2)

            # Série temporelle
            fig_temp = graphique_temp_horaire(variantes, zone, seuil_t1, seuil_t2)
            img = _fig_to_image(fig_temp)
            doc.add_picture(img, width=Cm(17))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # T_op vs T_ext (première variante)
            fig_op = graphique_text_vs_text_op(variantes[0], zone)
            img2 = _fig_to_image(fig_op)
            doc.add_picture(img2, width=Cm(17))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Diagramme bioclimatique (Givoni / COCO) — conditions intérieures
            series = []
            for v in variantes:
                pts = v.points_interieurs_givoni(zone, config, methode)
                if len(pts['T']):
                    pts['label'] = v.nom
                    series.append(pts)
            if series:
                fig_giv = creer_givoni(series, config=config, methode=methode,
                                       note_chauffe=note_chauffe)
                img3 = _fig_to_image(fig_giv)
                doc.add_picture(img3, width=Cm(17))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Diagramme(s) par période de focus (module optionnel)
            if periodes_on and periodes:
                for v in variantes:
                    pts = v.points_interieurs_par_periode(zone, config, periodes, methode)
                    if not len(pts['T']):
                        continue
                    pts['label'] = v.nom
                    fig_per = creer_givoni([pts], config=config, methode=methode,
                                           titre=f"{zone} · {v.nom} — par période",
                                           par_periode=True, note_chauffe=note_chauffe)
                    doc.add_picture(_fig_to_image(fig_per), width=Cm(17))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Récap heures + % inconfort par période, à 0 et 1 m/s
                def _pct(x):
                    return f"{x:.1f} %" if x == x else 'NA'
                rows_per = []
                for v in variantes:
                    r0 = {r['periode']: r for r in v.inconfort_periodes(zone, config, periodes, methode, 0.0)}
                    r1 = {r['periode']: r for r in v.inconfort_periodes(zone, config, periodes, methode, 1.0)}
                    for p in periodes:
                        nom = p.get('nom', '')
                        a = r0.get(nom, {})
                        b = r1.get(nom, {})
                        rows_per.append({
                            'Variante': v.nom, 'Période': nom,
                            "Heures d'occ.": f"{a.get('heures_occ', 0):.0f}",
                            "Inconfort 0 m/s (h)": f"{a.get('heures_inconfort', 0):.0f}",
                            '% inconfort 0 m/s': _pct(a.get('pct', float('nan'))),
                            "Inconfort 1 m/s (h)": f"{b.get('heures_inconfort', 0):.0f}",
                            '% inconfort 1 m/s': _pct(b.get('pct', float('nan'))),
                        })
                if rows_per:
                    _table_df(doc, _pd.DataFrame(rows_per).set_index(['Variante', 'Période']))

            # Apports solaires + internes
            fig_sol = graphique_apports_solaires(variantes, zone, type_apport="solaires")
            doc.add_picture(_fig_to_image(fig_sol), width=Cm(17))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_int = graphique_apports_solaires(variantes, zone, type_apport="internes")
            doc.add_picture(_fig_to_image(fig_int), width=Cm(17))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Carte(s) de température jour × heure (échelle commune aux variantes)
            tmins = [var.stats_temp(zone)['t_min'] for var in variantes]
            tmaxs = [var.stats_temp(zone)['t_max'] for var in variantes]
            tmins = [t for t in tmins if t == t]
            tmaxs = [t for t in tmaxs if t == t]
            zmin = min(tmins) if tmins else None
            zmax = max(tmaxs) if tmaxs else None
            for var in variantes:
                fig_thm = heatmap_temp_jour_heure(var, zone, zmin=zmin, zmax=zmax)
                if fig_thm is not None:
                    doc.add_picture(_fig_to_image(fig_thm), width=Cm(17))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Humidité relative : moyenne journalière (comparaison) + cartes jour×heure
            if any(not v.col_hr(zone).empty for v in variantes):
                seuils_cfg = config.get('seuils', {}) if isinstance(config, dict) else {}
                hr_min = float(seuils_cfg.get('hr_confort_min', 40.0))
                hr_max = float(seuils_cfg.get('hr_confort_max', 70.0))
                fig_hr = graphique_hr_horaire(variantes, zone, hr_min=hr_min, hr_max=hr_max,
                                              agregation="journalier")
                doc.add_picture(_fig_to_image(fig_hr), width=Cm(17))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for var in variantes:
                    fig_hhm = heatmap_hr_jour_heure(var, zone)
                    if fig_hhm is not None:
                        doc.add_picture(_fig_to_image(fig_hhm), width=Cm(17))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()

    # -- 3. Comparaison zones --
    if zones_comparaison and variantes:
        _para_rouge(doc, "3. Comparaison de zones", level=1)
        var = variantes[0]

        fig_t = graphique_temp_min_moy_max([var], zones_comparaison)
        doc.add_picture(_fig_to_image(fig_t), width=Cm(17))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        fig_dep = graphique_heures_depassement([var], zones_comparaison, seuil_t1, seuil_t2)
        doc.add_picture(_fig_to_image(fig_dep), width=Cm(17))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        fig_sol = graphique_apports_par_zone_mensuel(var, zones_comparaison, type_apport="solaires")
        doc.add_picture(_fig_to_image(fig_sol), width=Cm(17))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Restaurer le mode d'affichage initial (le rapport est rendu en clair)
    _charte.set_dark_mode(_prev_dark)

    # Sauvegarder dans un buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
